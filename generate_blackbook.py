"""
Generate MUSIC_STREAMING_SYSTEM_BLACK_BOOK.docx directly using python-docx
Complete black book with all actual project source codes and explanations.
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# ── Page margins ──────────────────────────────────────────────────────────────
for s in doc.sections:
    s.top_margin    = Inches(1)
    s.bottom_margin = Inches(1)
    s.left_margin   = Inches(1.25)
    s.right_margin  = Inches(1.25)

# ── Default font ──────────────────────────────────────────────────────────────
doc.styles['Normal'].font.name = 'Calibri'
doc.styles['Normal'].font.size = Pt(11)

# ─────────────────────────────────────────────────────────────────────────────
# HELPERS
# ─────────────────────────────────────────────────────────────────────────────
NAVY   = RGBColor(0x1a, 0x1a, 0x6e)
BLUE   = RGBColor(0x1f, 0x4e, 0x79)
MBLUE  = RGBColor(0x2e, 0x74, 0xb5)
WHITE  = RGBColor(0xFF, 0xFF, 0xFF)
BLACK  = RGBColor(0x00, 0x00, 0x00)
GRAY   = RGBColor(0x44, 0x44, 0x44)
RED    = RGBColor(0xC7, 0x25, 0x4E)

def h(text, level=1):
    p = doc.add_heading(text, level=level)
    colors = {1: NAVY, 2: BLUE, 3: MBLUE, 4: MBLUE}
    sizes  = {1: 18,   2: 15,   3: 13,    4: 12}
    for run in p.runs:
        run.font.color.rgb = colors.get(level, BLACK)
        run.font.size      = Pt(sizes.get(level, 11))
        run.bold           = True
    return p

def para(text, size=11, bold=False, italic=False, color=None, align=None, space_after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_after  = Pt(space_after)
    p.paragraph_format.space_before = Pt(2)
    if align == 'center':
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.font.size   = Pt(size)
    r.bold        = bold
    r.italic      = italic
    if color:
        r.font.color.rgb = color
    return p

def bullet(text, size=10.5):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent  = Inches(0.3)
    p.paragraph_format.space_after  = Pt(2)
    r = p.add_run(text)
    r.font.size = Pt(size)
    return p

_num_counter = [0]  # mutable so inner function can modify

def numbered(text, size=10.5, restart=False):
    if restart:
        _num_counter[0] = 0
    _num_counter[0] += 1
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Inches(0.5)
    p.paragraph_format.first_line_indent = Inches(-0.25)
    p.paragraph_format.space_after  = Pt(2)
    r = p.add_run(f'{_num_counter[0]}.  {text}')
    r.font.size = Pt(size)
    return p

def img(path, caption='', width=5.5):
    try:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after  = Pt(4)
        run = p.add_run()
        run.add_picture(path, width=Inches(width))
    except Exception as e:
        para(f'[Image not found: {path}]', italic=True, color=GRAY)
    if caption:
        cp = doc.add_paragraph()
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cp.paragraph_format.space_after = Pt(10)
        r = cp.add_run(caption)
        r.font.size  = Pt(9)
        r.italic     = True
        r.font.color.rgb = GRAY

def numbered_reset():
    _num_counter[0] = 0

def hr():
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bot = OxmlElement('w:bottom')
    bot.set(qn('w:val'),   'single')
    bot.set(qn('w:sz'),    '6')
    bot.set(qn('w:space'), '1')
    bot.set(qn('w:color'), '1F4E79')
    pBdr.append(bot)
    pPr.append(pBdr)

def code(text, lang=''):
    if lang:
        lbl = doc.add_paragraph()
        lbl.paragraph_format.left_indent = Inches(0.3)
        lbl.paragraph_format.space_after = Pt(0)
        r = lbl.add_run(f'  {lang.upper()}')
        r.font.size  = Pt(8)
        r.bold       = True
        r.font.color.rgb = MBLUE

    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Inches(0.3)
    p.paragraph_format.right_indent = Inches(0.3)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(8)

    r = p.add_run(text)
    r.font.name  = 'Courier New'
    r.font.size  = Pt(8.5)
    r.font.color.rgb = RGBColor(0x1a, 0x1a, 0x1a)

def table(headers, rows, col_widths=None):
    t = doc.add_table(rows=1+len(rows), cols=len(headers))
    t.style = 'Table Grid'
    # header row
    hrow = t.rows[0]
    for j, h_text in enumerate(headers):
        cell = hrow.cells[j]
        cell.text = h_text
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after  = Pt(3)
        for run in p.runs:
            run.bold = True
            run.font.size = Pt(9)
            run.font.color.rgb = WHITE
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'),   'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'),  '1F4E79')
        tcPr.append(shd)
    # data rows
    for i, row_data in enumerate(rows):
        row = t.rows[i+1]
        for j, cell_text in enumerate(row_data):
            cell = row.cells[j]
            cell.text = str(cell_text)
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after  = Pt(2)
            for run in p.runs:
                run.font.size = Pt(9)
    doc.add_paragraph()

def pb():
    doc.add_page_break()

# ─────────────────────────────────────────────────────────────────────────────
# TITLE PAGE
# ─────────────────────────────────────────────────────────────────────────────
doc.add_paragraph()
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('MUSIC STREAMING SYSTEM')
r.font.size = Pt(26); r.bold = True; r.font.color.rgb = NAVY

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('A Full-Stack Web Application for Online Music Discovery and Playback')
r.font.size = Pt(14); r.italic = True; r.font.color.rgb = BLUE

doc.add_paragraph()
hr()
doc.add_paragraph()

for label, value in [
    ('Project Title',    'Music Streaming System'),
    ('Technology Stack', 'React.js | FastAPI (Python) | SQLite | YouTube Data API v3'),
    ('Academic Year',    '2024–2025'),
    ('Project Type',     'Major Project — Full Stack Web Application'),
]:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = p.add_run(f'{label}: '); r1.bold = True; r1.font.size = Pt(11)
    r2 = p.add_run(value);        r2.font.size = Pt(11)

pb()

# ─────────────────────────────────────────────────────────────────────────────
# CERTIFICATE
# ─────────────────────────────────────────────────────────────────────────────
h('CERTIFICATE', 1)
para('This is to certify that the project entitled "Music Streaming System" has been carried out as a part of the curriculum for the award of the degree of Bachelor of Engineering / Bachelor of Technology in Computer Science & Engineering / Information Technology.')
para('This project work is a bonafide record of the original work done under the guidance of the project supervisor and has not been submitted elsewhere for the award of any degree or diploma.')
para('The project demonstrates the complete design, development, and testing of a full-stack music streaming web application using React.js, FastAPI, SQLAlchemy ORM, SQLite, and the YouTube Data API v3.')
doc.add_paragraph()
para('Project Guide Signature: ___________________')
para('Head of Department Signature: ___________________')
para('Date: ___________________')
pb()

# ─────────────────────────────────────────────────────────────────────────────
# ABSTRACT
# ─────────────────────────────────────────────────────────────────────────────
h('ABSTRACT', 1)
para('The Music Streaming System is a full-stack web application that enables users to search, stream, and manage music online using the YouTube Data API v3 as the content source. The system is built with a React.js single-page application frontend and a FastAPI Python backend, connected to an SQLite relational database through SQLAlchemy ORM.')
para('The system implements JWT-based user authentication with bcrypt password hashing, YouTube music search with automatic rotation across 30 API keys providing 300,000 daily quota units, a 12-hour client-side search result cache that reduces API consumption by approximately 80%, complete playlist management with CRUD operations, listening history tracking, mood-based music discovery, offline song bookmarking, social sharing across 6 platforms, voice search using the Web Speech API, karaoke mode with synchronized lyrics, a comprehensive analytics dashboard, and a persistent mini player with queue management.')
para('Keywords: Music Streaming, React.js, FastAPI, YouTube API, JWT Authentication, REST API, SQLAlchemy, SQLite, Web Speech API, Search Cache', italic=True)
pb()

# ─────────────────────────────────────────────────────────────────────────────
# ACKNOWLEDGEMENT
# ─────────────────────────────────────────────────────────────────────────────
h('ACKNOWLEDGEMENT', 1)
para('We take this opportunity to express our sincere gratitude to all those who have contributed to the successful completion of this project.')
para('We are deeply grateful to our project guide for their invaluable guidance, constant encouragement, and constructive suggestions throughout the development of this project.')
para('We also acknowledge the open-source communities behind FastAPI, React.js, SQLAlchemy, and the YouTube Data API documentation, whose tools and documentation made this project possible.')
pb()

# ─────────────────────────────────────────────────────────────────────────────
# TABLE OF CONTENTS
# ─────────────────────────────────────────────────────────────────────────────
h('TABLE OF CONTENTS', 1)
table(
    ['Chapter', 'Title'],
    [
        ['1',   'Introduction'],
        ['1.1', 'Background'],
        ['1.2', 'Objectives'],
        ['1.3', 'Purpose and Scope'],
        ['1.4', 'Achievements'],
        ['1.5', 'Organisation of Report'],
        ['2',   'Survey of Technologies'],
        ['3',   'Requirements and Analysis'],
        ['4',   'System Design'],
        ['4.1', 'Database Schema Design'],
        ['4.2', 'Authentication Design'],
        ['4.3', 'API Key Rotation Design'],
        ['4.4', 'Search Cache Design'],
        ['4.5', 'Procedural Design — Algorithms & Data Structures'],
        ['4.6', 'User Interface Design'],
        ['4.7', 'Security Design'],
        ['4.8', 'Test Case Design'],
        ['5',   'Implementation and Testing'],
        ['5.1', 'Backend Implementation (with Source Code)'],
        ['5.2', 'Frontend Implementation (with Source Code)'],
        ['5.3', 'Testing Results'],
        ['6',   'Results and Discussion'],
        ['7',   'Conclusions and Future Scope'],
        ['',    'References'],
        ['',    'Glossary'],
        ['A',   'Appendix A: Project Setup Guide'],
        ['B',   'Appendix B: API Endpoints Reference'],
    ]
)
pb()

# ─────────────────────────────────────────────────────────────────────────────
# CHAPTER 1
# ─────────────────────────────────────────────────────────────────────────────
h('CHAPTER 1: INTRODUCTION', 1)

h('1.1 Background', 2)
para('The digital transformation of the music industry has fundamentally changed how people discover, access, and enjoy music. With the proliferation of high-speed internet and smart devices, online music streaming has become the dominant mode of music consumption worldwide. Platforms such as Spotify, Apple Music, and YouTube Music have demonstrated the immense demand for feature-rich, personalized music streaming experiences.')
para('The Music Streaming System was developed to address these challenges in an academic context. The system leverages the YouTube Data API v3 as its content source — providing access to millions of music videos — while building all user management, playlist, history, analytics, and discovery features from scratch using modern web technologies. The backend is built with FastAPI connected to SQLite via SQLAlchemy ORM. The frontend is a React.js SPA communicating via RESTful API calls.')

h('1.2 Objectives', 2)
numbered_reset()
for obj in [
    'Design and develop a secure, multi-user music streaming web application',
    'Implement JWT-based user authentication with bcrypt password hashing',
    'Integrate YouTube Data API v3 for music search, streaming, and content discovery',
    'Build an intelligent API key rotation system across 30 keys (300,000 daily quota units)',
    'Develop a 12-hour client-side search result cache to reduce API consumption by up to 80%',
    'Implement complete playlist management with create, read, update, and delete operations',
    'Track user listening history and generate personalized music recommendations',
    'Build a mood-based music discovery feature using an interactive slider interface',
    'Implement voice search using the Web Speech API with NLP command processing',
    'Enable social sharing of songs across 6 major platforms',
    'Develop a comprehensive analytics dashboard showing user listening statistics',
    'Create a persistent mini player with queue management and next/previous navigation',
]:
    numbered(obj)

h('1.3 Purpose and Scope', 2)
para('Purpose: The Music Streaming System serves as a demonstration of full-stack web development capabilities, integrating modern frontend and backend technologies to deliver a production-quality music streaming experience.')
para('Scope covers: User registration/login/JWT management, YouTube search with language filtering (Hindi, English, Punjabi, Marathi, Tamil, Telugu), playlist CRUD, listening history, like/unlike songs, offline bookmarking, trending songs, mood-based discovery, analytics, voice search, social sharing, karaoke mode, mini player, dark/light theme, and onboarding flow.')

h('1.4 Achievements', 2)
numbered_reset()
for ach in [
    '30-key YouTube API rotation system providing 300,000 daily quota units with automatic failover',
    '12-hour localStorage search cache reducing API consumption by ~80% for repeat searches',
    'Complete JWT authentication with bcrypt — passwords never stored in plain text',
    '3-strategy recommendation engine: liked artists → trending → database fallback',
    'Voice search with NLP command processing (strips "play", "search for", "find" prefixes)',
    'Social sharing on 6 platforms: WhatsApp, Facebook, Twitter, Telegram, LinkedIn, Reddit',
    'Mood-based discovery with debounced API calls to prevent excessive requests',
    'Analytics dashboard: top songs, top artists, total plays, playlists, likes, listening time',
    'Karaoke mode with synchronized lyrics and auto-scroll',
    'Responsive UI with dark/light theme and 6-step onboarding flow for new users',
]:
    numbered(ach)

h('1.5 Organisation of Report', 2)
for item in [
    'Chapter 1 — Introduction: background, objectives, scope, achievements',
    'Chapter 2 — Survey of Technologies: all frameworks, libraries, and APIs used',
    'Chapter 3 — Requirements and Analysis: functional/non-functional requirements, architecture',
    'Chapter 4 — System Design: database schema, authentication, API design, algorithms, security, test cases',
    'Chapter 5 — Implementation and Testing: complete source code with explanations and test results',
    'Chapter 6 — Results and Discussion: performance metrics, feature completeness, challenges',
    'Chapter 7 — Conclusions: summary and future scope',
    'Appendices — Setup guide, API reference',
]:
    bullet(item)
pb()

# ─────────────────────────────────────────────────────────────────────────────
# CHAPTER 2
# ─────────────────────────────────────────────────────────────────────────────
h('CHAPTER 2: SURVEY OF TECHNOLOGIES', 1)

h('2.1 Frontend Technologies', 2)
h('2.1.1 React.js (v18)', 3)
para('React.js is an open-source JavaScript library developed by Meta for building user interfaces using a component-based architecture. React uses a Virtual DOM to efficiently update only changed parts of the actual DOM. In the Music Streaming System, React manages all UI components. Key hooks used: useState (local state), useEffect (API calls, timers), useRef (DOM access for lyrics scroll and speech recognition), useParams (URL parameters), useLocation (active route detection), useSearchParams (URL query parameters).')

h('2.1.2 React Router DOM (v6)', 3)
para('Provides client-side routing for the SPA. BrowserRouter wraps the application, Routes/Route define page mappings, Navigate handles redirects. Enables navigation between Home, Search, Playlists, History, Trending, Analytics, Offline, and Mood pages without full page reloads.')

h('2.1.3 Axios', 3)
para('Promise-based HTTP client. A custom Axios instance is configured with the backend base URL and an interceptor that automatically attaches the JWT Bearer token from localStorage to every outgoing request.')

h('2.1.4 Web Speech API', 3)
para('Browser-native JavaScript API for speech recognition. The SpeechRecognition interface (with webkitSpeechRecognition fallback) captures microphone input and returns both interim (real-time) and final transcripts. Used for the voice search feature.')

h('2.2 Backend Technologies', 2)
h('2.2.1 FastAPI (Python)', 3)
para('High-performance Python web framework built on Starlette and Pydantic. Provides automatic OpenAPI documentation, request/response validation, dependency injection, and native async/await support. Routes are organized into router modules: auth, youtube, playlists, history, stats, offline.')

h('2.2.2 SQLAlchemy ORM', 3)
para('Python SQL toolkit and Object Relational Mapper. Database tables are defined as Python classes (models). SQLAlchemy session management, relationship definitions, and query API are used throughout the backend.')

h('2.2.3 SQLite', 3)
para('Lightweight file-based relational database. Stores the entire database in music_sagar.db. Ideal for development. SQLAlchemy\'s check_same_thread=False enables concurrent request handling with FastAPI.')

h('2.2.4 Python Jose (JWT)', 3)
para('Provides JSON Web Token creation and verification. Tokens are signed with HS256 algorithm, contain the user ID as the sub claim, and expire after 7 days. Stateless authentication — no server-side session storage needed.')

h('2.2.5 Passlib + Bcrypt', 3)
para('Passlib CryptContext with bcrypt scheme provides secure one-way password hashing. Bcrypt is intentionally slow to prevent brute-force attacks. verify() uses constant-time comparison to prevent timing attacks.')

h('2.2.6 HTTPX', 3)
para('Async HTTP client for Python. Used in the YouTube router for non-blocking requests to YouTube Data API v3. Configurable timeouts: 5 seconds for search, 3 seconds for duration fetch.')

h('2.3 YouTube Data API v3', 2)
para('Two endpoints are used:')
bullet('Search endpoint (/youtube/v3/search): searches music videos, filtered to category 10 (Music), returns 20 results, supports pagination via pageToken. Costs 100 quota units per request.')
bullet('Videos endpoint (/youtube/v3/videos): fetches video duration in ISO 8601 format (e.g., PT3M45S), parsed to readable format (3:45). Costs 1 unit per video.')
para('Quota: Each key has 10,000 units/day. 30 keys = 300,000 units/day. 12-hour cache reduces consumption by ~80%.')

h('2.4 Development Tools', 2)
table(
    ['Tool', 'Purpose'],
    [
        ['Vite',          'Fast frontend build tool and dev server for React'],
        ['python-dotenv', 'Loads environment variables from .env files'],
        ['CORS Middleware','FastAPI middleware for cross-origin requests from React'],
        ['localStorage',  'Browser storage for JWT, cache, recent searches, onboarding state'],
        ['uvicorn',       'ASGI server for running the FastAPI application'],
        ['python-docx',   'Python library for generating Word (.docx) documents'],
    ]
)
pb()

# ─────────────────────────────────────────────────────────────────────────────
# CHAPTER 3
# ─────────────────────────────────────────────────────────────────────────────
h('CHAPTER 3: REQUIREMENTS AND ANALYSIS', 1)

h('3.1 Functional Requirements', 2)
table(
    ['ID', 'Requirement', 'Priority', 'Module'],
    [
        ['FR1',  'User registration with username, email, password',                    'High',   'Auth'],
        ['FR2',  'User login with JWT token generation (7-day expiry)',                 'High',   'Auth'],
        ['FR3',  'User logout with client-side token removal',                          'High',   'Auth'],
        ['FR4',  'Search YouTube for music by query string',                            'High',   'Search'],
        ['FR5',  'Filter search by language (Hindi, English, Punjabi, Marathi, Tamil)', 'High',   'Search'],
        ['FR6',  'Load more results via pagination (pageToken)',                        'Medium', 'Search'],
        ['FR7',  'Cache search results for 12 hours to reduce API calls',               'High',   'Search'],
        ['FR8',  'Voice search using microphone with NLP command processing',           'Medium', 'Search'],
        ['FR9',  'Play music via YouTube embedded player',                              'High',   'Player'],
        ['FR10', 'Mini player with play/pause, next, previous, progress bar',           'High',   'Player'],
        ['FR11', 'Full-screen player with karaoke mode',                                'Medium', 'Player'],
        ['FR12', 'Queue management — add, remove, clear songs',                         'Medium', 'Player'],
        ['FR13', 'Create, rename, and delete playlists',                                'High',   'Playlists'],
        ['FR14', 'Add songs to playlists from search results',                          'High',   'Playlists'],
        ['FR15', 'View all songs in a playlist',                                        'High',   'Playlists'],
        ['FR16', 'Automatic listening history tracking',                                'Medium', 'History'],
        ['FR17', 'View listening history (last 50 songs)',                              'Medium', 'History'],
        ['FR18', 'Clear entire listening history',                                      'Medium', 'History'],
        ['FR19', 'Like and unlike songs',                                               'Medium', 'Interactions'],
        ['FR20', 'Save songs for offline bookmarking',                                  'Medium', 'Offline'],
        ['FR21', 'View trending songs by play count',                                   'Medium', 'Discovery'],
        ['FR22', 'Mood-based music discovery (0–100 slider)',                           'Medium', 'Discovery'],
        ['FR23', 'Personalized song recommendations',                                   'Medium', 'Discovery'],
        ['FR24', 'Analytics dashboard (plays, playlists, likes, time)',                 'Low',    'Analytics'],
        ['FR25', 'Social sharing on 6 platforms + copy link',                           'Low',    'Sharing'],
        ['FR26', 'Dark/light theme toggle',                                             'Low',    'UI'],
        ['FR27', 'Onboarding flow for new users (6 steps)',                             'Low',    'UI'],
    ]
)

h('3.2 Non-Functional Requirements', 2)
table(
    ['ID', 'Requirement', 'Metric'],
    [
        ['NFR1', 'Search results load within 5 seconds',              'Response time'],
        ['NFR2', 'Cached results load within 100ms',                  'Response time'],
        ['NFR3', 'Passwords hashed using bcrypt',                     'Security'],
        ['NFR4', 'All protected routes require valid JWT',            'Security'],
        ['NFR5', 'System handles 30 API key rotation automatically',  'Availability'],
        ['NFR6', 'Cache supports up to 200 entries',                  'Scalability'],
        ['NFR7', 'UI responsive on desktop and tablet',               'Usability'],
        ['NFR8', 'Database queries complete within 200ms',            'Performance'],
    ]
)

h('3.3 System Architecture', 2)
para('The Music Streaming System follows a three-tier client-server architecture:')
bullet('Tier 1 — Presentation Layer: React.js SPA in the browser. Handles UI rendering, user interactions, client-side routing, and local state. Communicates with backend via HTTP REST API using Axios.')
bullet('Tier 2 — Application Layer: FastAPI Python application on uvicorn ASGI server. Handles business logic, authentication, database operations, and YouTube API integration. Organized into router modules.')
bullet('Tier 3 — Data Layer: SQLite database via SQLAlchemy ORM. Stores users, songs, playlists, history, likes, offline songs. Song records are shared across all users to prevent duplication.')
bullet('External Service: YouTube Data API v3 provides music video search results and metadata. Accessed from the backend to keep API keys secure.')
img(r'C:\Users\hp\Documents\a\b\System_Architecture.png', 'Figure 3.1: Three-Tier System Architecture', width=6.0)

pb()

# ─────────────────────────────────────────────────────────────────────────────
# CHAPTER 4
# ─────────────────────────────────────────────────────────────────────────────
h('CHAPTER 4: SYSTEM DESIGN', 1)

h('4.1 Database Schema Design', 2)
h('4.1.1 Database Tables and Attributes', 3)
table(['Table', 'Field', 'Data Type', 'Constraints', 'Description'], [
    ['users',         'id',               'INTEGER',  'PK, Auto Increment',        'Unique user identifier'],
    ['users',         'username',         'VARCHAR',  'Unique, Not Null, Indexed', 'Display name'],
    ['users',         'email',            'VARCHAR',  'Unique, Not Null, Indexed', 'Email address'],
    ['users',         'hashed_password',  'VARCHAR',  'Not Null',                  'bcrypt-hashed password'],
    ['users',         'created_at',       'DATETIME', 'Default: utcnow',           'Account creation time'],
    ['users',         'theme',            'VARCHAR',  'Default: dark',             'UI theme preference'],
    ['songs',         'id',               'INTEGER',  'PK, Auto Increment',        'Unique song identifier'],
    ['songs',         'youtube_video_id', 'VARCHAR',  'Unique, Not Null, Indexed', 'YouTube video ID'],
    ['songs',         'title',            'VARCHAR',  'Not Null',                  'Song title'],
    ['songs',         'thumbnail',        'VARCHAR',  'Nullable',                  'YouTube thumbnail URL'],
    ['songs',         'channel',          'VARCHAR',  'Nullable',                  'YouTube channel name'],
    ['playlists',     'id',               'INTEGER',  'PK, Auto Increment',        'Unique playlist ID'],
    ['playlists',     'name',             'VARCHAR',  'Not Null',                  'Playlist name'],
    ['playlists',     'user_id',          'INTEGER',  'FK → users.id',             'Owner of playlist'],
    ['playlist_songs','playlist_id',      'INTEGER',  'FK → playlists.id',         'Reference to playlist'],
    ['playlist_songs','song_id',          'INTEGER',  'FK → songs.id',             'Reference to song'],
    ['history',       'user_id',          'INTEGER',  'FK → users.id',             'User who played'],
    ['history',       'song_id',          'INTEGER',  'FK → songs.id',             'Song that was played'],
    ['history',       'played_at',        'DATETIME', 'Default: utcnow',           'Playback timestamp'],
    ['likes',         'user_id',          'INTEGER',  'FK → users.id',             'User who liked'],
    ['likes',         'song_id',          'INTEGER',  'FK → songs.id',             'Song that was liked'],
    ['likes',         'created_at',       'DATETIME', 'Default: utcnow',           'Like timestamp'],
    ['offline_songs', 'user_id',          'INTEGER',  'FK → users.id',             'User who saved'],
    ['offline_songs', 'song_id',          'INTEGER',  'FK → songs.id',             'Song that was saved'],
    ['offline_songs', 'saved_at',         'DATETIME', 'Default: utcnow',           'Save timestamp'],
])
img(r'C:\Users\hp\Documents\a\b\ER_Diagram.png', 'Figure 4.1: Entity Relationship Diagram', width=6.2)

h('4.1.2 SQLAlchemy ORM Models — backend/app/models.py', 3)
code('''from sqlalchemy import Column, Integer, String, DateTime, ForeignKey
from sqlalchemy.orm import relationship
from datetime import datetime
from .database import Base

class User(Base):
    __tablename__ = "users"
    id              = Column(Integer, primary_key=True, index=True)
    username        = Column(String, unique=True, index=True, nullable=False)
    email           = Column(String, unique=True, index=True, nullable=False)
    hashed_password = Column(String, nullable=False)
    created_at      = Column(DateTime, default=datetime.utcnow)
    theme           = Column(String, default="dark")
    playlists    = relationship("Playlist",    back_populates="user")
    history      = relationship("History",     back_populates="user")
    likes        = relationship("Like",        back_populates="user")
    offline_songs= relationship("OfflineSong", back_populates="user")

class Song(Base):
    __tablename__ = "songs"
    id               = Column(Integer, primary_key=True, index=True)
    youtube_video_id = Column(String, unique=True, index=True, nullable=False)
    title            = Column(String, nullable=False)
    thumbnail        = Column(String)
    channel          = Column(String)
    playlist_songs = relationship("PlaylistSong", back_populates="song")
    history        = relationship("History",      back_populates="song")
    likes          = relationship("Like",         back_populates="song")
    offline_songs  = relationship("OfflineSong",  back_populates="song")

class Playlist(Base):
    __tablename__ = "playlists"
    id      = Column(Integer, primary_key=True, index=True)
    name    = Column(String, nullable=False)
    user_id = Column(Integer, ForeignKey("users.id"), nullable=False)
    user          = relationship("User",        back_populates="playlists")
    playlist_songs= relationship("PlaylistSong",back_populates="playlist")

class PlaylistSong(Base):
    __tablename__ = "playlist_songs"
    id          = Column(Integer, primary_key=True, index=True)
    playlist_id = Column(Integer, ForeignKey("playlists.id"), nullable=False)
    song_id     = Column(Integer, ForeignKey("songs.id"),     nullable=False)
    playlist = relationship("Playlist", back_populates="playlist_songs")
    song     = relationship("Song",     back_populates="playlist_songs")

class History(Base):
    __tablename__ = "history"
    id        = Column(Integer, primary_key=True, index=True)
    user_id   = Column(Integer, ForeignKey("users.id"), nullable=False)
    song_id   = Column(Integer, ForeignKey("songs.id"), nullable=False)
    played_at = Column(DateTime, default=datetime.utcnow)
    user = relationship("User", back_populates="history")
    song = relationship("Song", back_populates="history")

class Like(Base):
    __tablename__ = "likes"
    id         = Column(Integer, primary_key=True, index=True)
    user_id    = Column(Integer, ForeignKey("users.id"), nullable=False)
    song_id    = Column(Integer, ForeignKey("songs.id"), nullable=False)
    created_at = Column(DateTime, default=datetime.utcnow)
    user = relationship("User", back_populates="likes")
    song = relationship("Song", back_populates="likes")

class OfflineSong(Base):
    __tablename__ = "offline_songs"
    id       = Column(Integer, primary_key=True, index=True)
    user_id  = Column(Integer, ForeignKey("users.id"), nullable=False)
    song_id  = Column(Integer, ForeignKey("songs.id"), nullable=False)
    saved_at = Column(DateTime, default=datetime.utcnow)
    user = relationship("User", back_populates="offline_songs")
    song = relationship("Song", back_populates="offline_songs")''', 'python')
para('The Song table uses youtube_video_id as a unique index. All features use a "get or create" pattern — if a song already exists in the songs table, the existing record is reused. This prevents duplicate entries and ensures play counts in the history table accurately reflect total plays across all users. The PlaylistSong junction table implements the many-to-many relationship between playlists and songs.')

h('4.1.3 Database Connection — backend/app/database.py', 3)
code('''import os
from sqlalchemy import create_engine
from sqlalchemy.ext.declarative import declarative_base
from sqlalchemy.orm import sessionmaker

DATABASE_URL = os.getenv("DATABASE_URL", "sqlite:///./music_sagar.db")

if DATABASE_URL.startswith("sqlite"):
    engine = create_engine(DATABASE_URL, connect_args={"check_same_thread": False})
else:
    engine = create_engine(DATABASE_URL)

SessionLocal = sessionmaker(autocommit=False, autoflush=False, bind=engine)
Base = declarative_base()

def get_db():
    """FastAPI dependency: provides a DB session per request, always closes after."""
    db = SessionLocal()
    try:
        yield db
    finally:
        db.close()''', 'python')

h('4.1.4 Pydantic Schemas — backend/app/schemas.py', 3)
code('''from pydantic import BaseModel, EmailStr
from datetime import datetime
from typing import Optional

class UserCreate(BaseModel):
    username: str
    email: EmailStr   # Validates email format automatically
    password: str

class UserLogin(BaseModel):
    email: EmailStr
    password: str

class UserResponse(BaseModel):
    id: int; username: str; email: str; created_at: datetime
    class Config: from_attributes = True

class SongBase(BaseModel):
    youtube_video_id: str; title: str
    thumbnail: Optional[str] = None; channel: Optional[str] = None

class SongResponse(SongBase):
    id: int
    class Config: from_attributes = True

class PlaylistCreate(BaseModel):
    name: str

class PlaylistResponse(BaseModel):
    id: int; name: str; user_id: int
    class Config: from_attributes = True

class AddSongToPlaylist(BaseModel):
    youtube_video_id: str; title: str
    thumbnail: Optional[str] = None; channel: Optional[str] = None

class HistoryCreate(BaseModel):
    youtube_video_id: str; title: str
    thumbnail: Optional[str] = None; channel: Optional[str] = None

class HistoryResponse(BaseModel):
    id: int; played_at: datetime; song: SongResponse
    class Config: from_attributes = True''', 'python')

h('4.2 Authentication Design — backend/app/auth.py', 2)
code('''import os
from datetime import datetime, timedelta
from typing import Optional
from jose import JWTError, jwt
from passlib.context import CryptContext
from fastapi import Depends, HTTPException, status
from fastapi.security import HTTPBearer, HTTPAuthorizationCredentials
from sqlalchemy.orm import Session
from .database import get_db
from .models import User

SECRET_KEY = os.getenv("SECRET_KEY", "your-secret-key-min-32-characters-long")
ALGORITHM  = "HS256"
ACCESS_TOKEN_EXPIRE_DAYS = 7

pwd_context = CryptContext(schemes=["bcrypt"], deprecated="auto")
security    = HTTPBearer()

def verify_password(plain_password: str, hashed_password: str) -> bool:
    return pwd_context.verify(plain_password, hashed_password)

def get_password_hash(password: str) -> str:
    return pwd_context.hash(password)

def create_access_token(data: dict, expires_delta: Optional[timedelta] = None) -> str:
    to_encode = data.copy()
    expire = datetime.utcnow() + (expires_delta or timedelta(days=ACCESS_TOKEN_EXPIRE_DAYS))
    to_encode.update({"exp": expire})
    return jwt.encode(to_encode, SECRET_KEY, algorithm=ALGORITHM)

def get_current_user(
    credentials: HTTPAuthorizationCredentials = Depends(security),
    db: Session = Depends(get_db)
) -> User:
    """FastAPI dependency for protected routes. Validates JWT and returns User."""
    exc = HTTPException(status_code=401, detail="Could not validate credentials",
                        headers={"WWW-Authenticate": "Bearer"})
    try:
        payload = jwt.decode(credentials.credentials, SECRET_KEY, algorithms=[ALGORITHM])
        user_id = payload.get("sub")
        if user_id is None: raise exc
    except JWTError:
        raise exc
    user = db.query(User).filter(User.id == user_id).first()
    if user is None: raise exc
    return user''', 'python')
para('The authentication system uses stateless JWT tokens. On login, the server creates a JWT containing the user ID as the sub claim, signed with HS256. For every protected request, the client sends the token in the Authorization: Bearer header. get_current_user() decodes and verifies the token, then fetches the user from the database. bcrypt uses constant-time comparison to prevent timing attacks.')
img(r'C:\Users\hp\Documents\a\b\Simple Flow Representation.png', 'Figure 4.1: JWT Authentication Flow', width=4.5)

h('4.3 API Key Rotation Design — backend/app/routers/youtube.py', 2)
code('''import os, httpx, asyncio, re
from fastapi import APIRouter, HTTPException, status
from typing import List, Optional

router = APIRouter(prefix="/youtube", tags=["YouTube"])
YOUTUBE_SEARCH_URL = "https://www.googleapis.com/youtube/v3/search"
YOUTUBE_VIDEOS_URL = "https://www.googleapis.com/youtube/v3/videos"

failed_keys = set()  # Module-level: tracks quota-exhausted keys

def get_youtube_api_key():
    """Iterate keys 1-30, skip exhausted ones. Reset if all fail."""
    for i in range(1, 31):
        key_name = f"YOUTUBE_API_KEY_{i}"
        api_key  = os.getenv(key_name)
        if api_key and key_name not in failed_keys:
            return api_key, key_name
    failed_keys.clear()
    return os.getenv("YOUTUBE_API_KEY_1"), "YOUTUBE_API_KEY_1"

def mark_key_as_failed(key_name: str):
    failed_keys.add(key_name)
    print(f"Key {key_name} exhausted. Failed: {len(failed_keys)}/30")

def parse_iso8601_duration(duration: str) -> Optional[str]:
    """Convert PT3M45S -> 3:45, PT1H2M30S -> 1:02:30"""
    if not duration: return None
    match = re.match(r\'PT(?:(\\d+)H)?(?:(\\d+)M)?(?:(\\d+)S)?\', duration)
    if not match: return None
    h = int(match.group(1) or 0)
    m = int(match.group(2) or 0)
    s = int(match.group(3) or 0)
    return f"{h}:{m:02d}:{s:02d}" if h > 0 else f"{m}:{s:02d}"

async def fetch_video_durations(video_ids: List[str], api_key: str) -> dict:
    """Fetch durations with 3-second timeout to avoid blocking search results."""
    if not video_ids: return {}
    try:
        params = {"part": "contentDetails", "id": ",".join(video_ids[:50]), "key": api_key}
        async with httpx.AsyncClient(timeout=3.0) as client:
            response = await client.get(YOUTUBE_VIDEOS_URL, params=params)
            response.raise_for_status()
        return {item["id"]: parse_iso8601_duration(
                    item.get("contentDetails", {}).get("duration", ""))
                for item in response.json().get("items", [])
                if parse_iso8601_duration(item.get("contentDetails", {}).get("duration", ""))}
    except Exception: return {}

@router.get("/search")
async def search_youtube(query: str, language: str = "", pageToken: str = None):
    """Search YouTube music with language filter, pagination, and key rotation."""
    YOUTUBE_API_KEY, key_name = get_youtube_api_key()
    if not YOUTUBE_API_KEY:
        raise HTTPException(500, "All YouTube API keys have exceeded their quota.")

    lang_map = {"hindi":"hindi","english":"english","punjabi":"punjabi",
                "marathi":"marathi","tamil":"tamil","telugu":"telugu"}
    search_query = f"{query} {lang_map[language.lower()]}" if language.lower() in lang_map else query

    params = {"part":"snippet","q":search_query,"type":"video",
              "videoCategoryId":"10","maxResults":20,"key":YOUTUBE_API_KEY}
    if pageToken: params["pageToken"] = pageToken

    async with httpx.AsyncClient(timeout=5.0) as client:
        response = await client.get(YOUTUBE_SEARCH_URL, params=params)
        if response.status_code == 403:
            reason = response.json().get(\'error\',{}).get(\'errors\',[{}])[0].get(\'reason\',\'\')
            if \'quota\' in reason.lower():
                mark_key_as_failed(key_name)
                YOUTUBE_API_KEY, key_name = get_youtube_api_key()
                params["key"] = YOUTUBE_API_KEY
                response = await client.get(YOUTUBE_SEARCH_URL, params=params)
        response.raise_for_status()
        data = response.json()

    results = [{"videoId": item["id"]["videoId"],
                "title": item["snippet"]["title"],
                "thumbnail": item["snippet"]["thumbnails"]["high"]["url"],
                "channelTitle": item["snippet"]["channelTitle"],
                "duration": None}
               for item in data.get("items", [])]
    video_ids = [r["videoId"] for r in results]

    try:
        durations = await asyncio.wait_for(
            fetch_video_durations(video_ids, YOUTUBE_API_KEY), timeout=3.0)
        for r in results:
            r["duration"] = durations.get(r["videoId"])
    except asyncio.TimeoutError:
        pass

    return {"results": results,
            "nextPageToken": data.get("nextPageToken"),
            "prevPageToken": data.get("prevPageToken")}''', 'python')
para('The failed_keys set is module-level, persisting across requests within the same server session. When a 403 quota error is received, the key is marked failed and the request is immediately retried with the next available key — providing seamless failover. Duration fetching runs as a separate async operation with a 3-second timeout so search results are never blocked.')
img(r'C:\Users\hp\Documents\a\b\API key diagram.png', 'Figure 4.2: API Key Rotation Logic', width=4.5)

h('4.4 Search Cache Design — frontend/src/utils/searchCache.js', 2)
code('''const CACHE_DURATION = 12 * 60 * 60 * 1000; // 12 hours
const MAX_CACHE_SIZE = 200;

class SearchCache {
  constructor() {
    this.cache = new Map();
    this.loadFromLocalStorage();
  }

  generateKey(query, language = \'all\') {
    return `${query.toLowerCase().trim()}_${language}`;
  }

  get(query, language = \'all\') {
    const key    = this.generateKey(query, language);
    const cached = this.cache.get(key);
    if (!cached) return null;
    if (Date.now() - cached.timestamp > CACHE_DURATION) {
      this.cache.delete(key);
      this.saveToLocalStorage();
      return null;
    }
    console.log(`Cache HIT: "${query}" — API call saved!`);
    return cached.data;
  }

  set(query, language = \'all\', data) {
    const key = this.generateKey(query, language);
    if (this.cache.size >= MAX_CACHE_SIZE) {
      this.cache.delete(this.cache.keys().next().value); // evict oldest
    }
    this.cache.set(key, { data, timestamp: Date.now() });
    this.saveToLocalStorage();
  }

  clearExpired() {
    const now = Date.now();
    for (const [key, value] of this.cache.entries())
      if (now - value.timestamp > CACHE_DURATION) this.cache.delete(key);
    this.saveToLocalStorage();
  }

  saveToLocalStorage() {
    try {
      localStorage.setItem(\'music_sagar_search_cache\',
        JSON.stringify(Array.from(this.cache.entries())));
    } catch (e) {}
  }

  loadFromLocalStorage() {
    try {
      const stored = localStorage.getItem(\'music_sagar_search_cache\');
      if (stored) { this.cache = new Map(JSON.parse(stored)); this.clearExpired(); }
    } catch (e) { this.cache = new Map(); }
  }
}

export const searchCache = new SearchCache();
setInterval(() => searchCache.clearExpired(), 2 * 60 * 60 * 1000);''', 'javascript')
para('The cache uses a JavaScript Map with composite keys (query_language). Each entry stores data and a timestamp. On retrieval, if the entry is older than 12 hours it is deleted and null is returned, triggering a fresh API call. The cache is serialized to localStorage as a JSON array of [key, value] pairs, ensuring persistence across page refreshes. Maximum 200 entries with LRU-style eviction. Proactive cleanup runs every 2 hours.')
img(r'C:\Users\hp\Documents\a\b\Search with Cache Logic.png', 'Figure 4.3: Search Cache Flow', width=4.5)

h('4.5 Procedural Design', 2)
h('4.5.1 Key Algorithms', 3)
para('User Authentication Algorithm:')
numbered_reset()
for step in ['User submits email and password via POST /auth/login',
             'Backend queries users table by email',
             'If user not found → return 401 Unauthorized',
             'verify_password(plain, hashed) compares using bcrypt constant-time comparison',
             'If mismatch → return 401 Unauthorized',
             'create_access_token({"sub": user_id}) generates signed JWT with 7-day expiry',
             'Return JWT token + user info to client',
             'Client stores token in localStorage for all future requests']:
    numbered(step)

para('Recommendation Algorithm (3-Strategy Fallback):')
numbered_reset()
for step in ['Query liked songs for current user (limit 10)',
             'If liked songs exist → collect unique artist names from liked songs',
             'For each artist → query songs by that artist (excluding already-liked songs)',
             'If recommendations < 50 → add trending songs (most played globally via history count)',
             'If still < 50 → add any songs from database as fallback',
             'Shuffle results for variety → return first 50']:
    numbered(step)

h('4.5.2 Data Structures Used', 3)
table(['Structure', 'Where Used', 'Purpose'], [
    ['Array / List',    'Search results, history, playlists, recommendations', 'Ordered collections of songs'],
    ['Set',             'failed_keys (Python), offlineSongs/likedSongs (JS)',  'O(1) lookup, no duplicates'],
    ['Map',             'searchCache (JavaScript)',                             'Key-value cache with ordered insertion'],
    ['Queue (Array)',   'Play queue in App.jsx',                               'FIFO song playback order'],
    ['Stack (Array)',   'playHistory in App.jsx',                              'LIFO for previous song navigation'],
    ['Dictionary',      'API response formatting (Python)',                    'Key-value data serialization'],
    ['Hash Table',      'SQLAlchemy session cache',                            'Fast ORM object lookup'],
])

h('4.6 User Interface Design', 2)
para('Layout: Sidebar (left, fixed) + Main Content Area (right, scrollable) + Mini Player (bottom, fixed overlay).')
bullet('Sidebar: Logo, username, 8 navigation links, theme toggle, logout button')
bullet('Main Content: Page-specific content rendered by React Router (Home, Search, Playlists, History, Trending, Analytics, Offline, Mood)')
bullet('Mini Player: Appears when song is playing and full-screen player is closed. Shows thumbnail, title, artist, progress bar, prev/play/next controls')
bullet('Full-Screen Player: YouTube embedded player with karaoke mode, queue panel, social sharing')
bullet('Toast Notifications: Non-blocking feedback messages (top-right corner)')
bullet('Dark theme: Background #0a0a0a, cards #1a1a2e, accent gradient #667eea → #764ba2')
bullet('Light theme: Background #f5f5f5, cards #ffffff, same accent gradient')

h('4.7 Security Design', 2)
table(['Security Concern', 'Implementation'], [
    ['Password Storage',  'bcrypt hashing via passlib — plain text never stored or logged'],
    ['Authentication',    'JWT tokens signed with HS256, 7-day expiry, stored in localStorage'],
    ['Authorization',     'All protected routes use get_current_user dependency; users access only their own data'],
    ['Input Validation',  'Pydantic schemas validate all request bodies — invalid data returns 422 automatically'],
    ['API Key Security',  'YouTube API keys stored in backend/.env, never exposed to frontend'],
    ['CORS',              'CORS middleware configured; restrict to frontend domain in production'],
    ['Timing Attacks',    'bcrypt verify() uses constant-time comparison'],
])

h('4.8 Test Case Design', 2)
table(['Test ID', 'Test Case', 'Input', 'Expected Output'], [
    ['TC1',  'Valid signup',                    'New username, email, password',  '201 Created, user object'],
    ['TC2',  'Duplicate email signup',          'Existing email',                 '400 Bad Request'],
    ['TC3',  'Valid login',                     'Correct email + password',       '200 OK, JWT token'],
    ['TC4',  'Invalid password',                'Wrong password',                 '401 Unauthorized'],
    ['TC5',  'Protected route without token',   'No Authorization header',        '401 Unauthorized'],
    ['TC6',  'Search valid query',              '"Arijit Singh"',                 '20 YouTube results'],
    ['TC7',  'Search with language filter',     'query="songs", language="hindi"','Hindi songs results'],
    ['TC8',  'Repeat search within 12 hours',   'Same query',                     'Instant cache hit'],
    ['TC9',  'Create playlist',                 'name="My Playlist"',             '201 Created'],
    ['TC10', 'Add song to playlist',            'Valid song data',                'Song added, 200 OK'],
    ['TC11', 'Add duplicate song to playlist',  'Same song twice',                '400 Bad Request'],
    ['TC12', 'Delete playlist',                 'Valid playlist ID',              '204 No Content'],
    ['TC13', 'Access another user\'s playlist', 'Different user\'s playlist ID', '404 Not Found'],
    ['TC14', 'Like a song',                     'Valid video ID',                 'Song liked, 200 OK'],
    ['TC15', 'Save song offline',               'Valid song data',                'Song saved, 201 Created'],
    ['TC16', 'Clear history',                   'Authenticated user',             'History cleared, 200 OK'],
])
pb()

# ─────────────────────────────────────────────────────────────────────────────
# CHAPTER 5 — IMPLEMENTATION
# ─────────────────────────────────────────────────────────────────────────────
h('CHAPTER 5: IMPLEMENTATION AND TESTING', 1)

h('5.1 Backend Implementation', 2)

# ── main.py ──
h('5.1.1 Application Entry Point — backend/app/main.py', 3)
code('''from pathlib import Path
from dotenv import load_dotenv
from fastapi import FastAPI
from fastapi.middleware.cors import CORSMiddleware
from .database import engine, Base
from .routers import auth, youtube, playlists, history, stats, offline, ai, proxy
from .init_db import init_database

load_dotenv(dotenv_path=Path(__file__).parent.parent / \'.env\')
Base.metadata.create_all(bind=engine)
init_database()

app = FastAPI(title="Music Streaming System API", version="1.0.0")
app.add_middleware(CORSMiddleware, allow_origins=["*"],
                   allow_credentials=False, allow_methods=["*"], allow_headers=["*"])

app.include_router(auth.router)
app.include_router(youtube.router)
app.include_router(playlists.router)
app.include_router(history.router)
app.include_router(stats.router)
app.include_router(offline.router)
app.include_router(ai.router)
app.include_router(proxy.router)

@app.get("/")
def root(): return {"message": "Welcome to Music Streaming System API"}

@app.get("/health")
def health_check(): return {"status": "healthy"}''', 'python')
para('On startup, Base.metadata.create_all() creates all database tables if they do not exist. Each feature area has its own router module, keeping the codebase organized. CORS middleware allows all origins for development — restrict to the frontend domain in production.')

# ── auth router ──
h('5.1.2 Authentication Router — backend/app/routers/auth.py', 3)
code('''from fastapi import APIRouter, Depends, HTTPException, status
from sqlalchemy.orm import Session
from datetime import timedelta
from ..database import get_db
from ..models import User
from ..schemas import UserCreate, UserLogin, UserResponse
from ..auth import get_password_hash, verify_password, create_access_token, get_current_user, ACCESS_TOKEN_EXPIRE_DAYS

router = APIRouter(prefix="/auth", tags=["Authentication"])

@router.post("/signup", response_model=UserResponse, status_code=201)
def signup(user: UserCreate, db: Session = Depends(get_db)):
    existing = db.query(User).filter(
        (User.email == user.email) | (User.username == user.username)).first()
    if existing:
        raise HTTPException(400, "Username or email already registered")
    new_user = User(username=user.username, email=user.email,
                    hashed_password=get_password_hash(user.password))
    db.add(new_user); db.commit(); db.refresh(new_user)
    return new_user

@router.post("/login")
def login(user: UserLogin, db: Session = Depends(get_db)):
    db_user = db.query(User).filter(User.email == user.email).first()
    if not db_user or not verify_password(user.password, db_user.hashed_password):
        raise HTTPException(401, "Incorrect email or password")
    token = create_access_token({"sub": str(db_user.id)},
                                 timedelta(days=ACCESS_TOKEN_EXPIRE_DAYS))
    return {"access_token": token, "token_type": "bearer",
            "user": {"id": db_user.id, "username": db_user.username, "email": db_user.email}}

@router.post("/logout")
def logout(): return {"message": "Logout successful"}

@router.get("/me", response_model=UserResponse)
def get_me(current_user: User = Depends(get_current_user)): return current_user

@router.post("/theme")
def update_theme(theme: str, current_user: User = Depends(get_current_user),
                 db: Session = Depends(get_db)):
    if theme not in ["dark", "light"]:
        raise HTTPException(400, "Theme must be dark or light")
    current_user.theme = theme; db.commit()
    return {"message": "Theme updated", "theme": theme}''', 'python')
para('signup checks for duplicate username or email using an OR filter. The password is hashed before storage — plain text is never saved. login returns a JWT token plus basic user info. The frontend stores this token in localStorage. logout is client-side only — the client removes the token from localStorage.')

# ── playlists router ──
h('5.1.3 Playlist Router — backend/app/routers/playlists.py', 3)
code('''from fastapi import APIRouter, Depends, status, HTTPException
from sqlalchemy.orm import Session
from typing import List
from ..database import get_db
from ..models import Playlist, Song, PlaylistSong, User
from ..schemas import PlaylistCreate, PlaylistResponse, AddSongToPlaylist, SongResponse
from ..auth import get_current_user

router = APIRouter(prefix="/playlists", tags=["Playlists"])

@router.post("", response_model=PlaylistResponse, status_code=201)
def create_playlist(playlist: PlaylistCreate, current_user: User = Depends(get_current_user),
                    db: Session = Depends(get_db)):
    p = Playlist(name=playlist.name, user_id=current_user.id)
    db.add(p); db.commit(); db.refresh(p); return p

@router.get("", response_model=List[PlaylistResponse])
def get_playlists(current_user: User = Depends(get_current_user), db: Session = Depends(get_db)):
    return db.query(Playlist).filter(Playlist.user_id == current_user.id).all()

@router.get("/{playlist_id}/songs")
def get_playlist_songs(playlist_id: int, current_user: User = Depends(get_current_user),
                       db: Session = Depends(get_db)):
    playlist = db.query(Playlist).filter(Playlist.id == playlist_id,
                                          Playlist.user_id == current_user.id).first()
    if not playlist: raise HTTPException(404, "Playlist not found")
    songs = db.query(Song).join(PlaylistSong, Song.id == PlaylistSong.song_id).filter(
        PlaylistSong.playlist_id == playlist_id).all()
    return {"playlist": {"id": playlist.id, "name": playlist.name},
            "songs": [{"youtube_video_id": s.youtube_video_id, "title": s.title,
                        "thumbnail": s.thumbnail, "channel": s.channel} for s in songs]}

@router.post("/{playlist_id}/add-song", response_model=SongResponse)
def add_song_to_playlist(playlist_id: int, song_data: AddSongToPlaylist,
                          current_user: User = Depends(get_current_user),
                          db: Session = Depends(get_db)):
    playlist = db.query(Playlist).filter(Playlist.id == playlist_id,
                                          Playlist.user_id == current_user.id).first()
    if not playlist: raise HTTPException(404, "Playlist not found")
    song = db.query(Song).filter(Song.youtube_video_id == song_data.youtube_video_id).first()
    if not song:
        song = Song(youtube_video_id=song_data.youtube_video_id, title=song_data.title,
                    thumbnail=song_data.thumbnail, channel=song_data.channel)
        db.add(song); db.commit(); db.refresh(song)
    if db.query(PlaylistSong).filter(PlaylistSong.playlist_id == playlist_id,
                                      PlaylistSong.song_id == song.id).first():
        raise HTTPException(400, "Song already in playlist")
    db.add(PlaylistSong(playlist_id=playlist_id, song_id=song.id)); db.commit()
    return song

@router.delete("/{playlist_id}", status_code=204)
def delete_playlist(playlist_id: int, current_user: User = Depends(get_current_user),
                    db: Session = Depends(get_db)):
    p = db.query(Playlist).filter(Playlist.id == playlist_id,
                                   Playlist.user_id == current_user.id).first()
    if not p: raise HTTPException(404, "Playlist not found")
    db.delete(p); db.commit()

@router.put("/{playlist_id}", response_model=PlaylistResponse)
def rename_playlist(playlist_id: int, playlist: PlaylistCreate,
                    current_user: User = Depends(get_current_user),
                    db: Session = Depends(get_db)):
    p = db.query(Playlist).filter(Playlist.id == playlist_id,
                                   Playlist.user_id == current_user.id).first()
    if not p: raise HTTPException(404, "Playlist not found")
    p.name = playlist.name; db.commit(); db.refresh(p); return p''', 'python')
para('All playlist endpoints filter by current_user.id ensuring users can only access their own playlists. add_song_to_playlist uses "get or create" for songs — if the YouTube video already exists in the songs table it is reused, preventing duplicates. Duplicate song detection in a playlist queries the PlaylistSong junction table.')

# ── history router ──
h('5.1.4 History Router — backend/app/routers/history.py', 3)
code('''from fastapi import APIRouter, Depends, HTTPException, status
from sqlalchemy.orm import Session
from typing import List
from ..database import get_db
from ..models import Song, History, User
from ..schemas import HistoryCreate, HistoryResponse
from ..auth import get_current_user

router = APIRouter(prefix="/history", tags=["History"])

@router.post("", status_code=201)
def add_to_history(history_data: HistoryCreate, current_user: User = Depends(get_current_user),
                   db: Session = Depends(get_db)):
    song = db.query(Song).filter(Song.youtube_video_id == history_data.youtube_video_id).first()
    if not song:
        song = Song(youtube_video_id=history_data.youtube_video_id, title=history_data.title,
                    thumbnail=history_data.thumbnail, channel=history_data.channel)
        db.add(song); db.commit(); db.refresh(song)
    db.add(History(user_id=current_user.id, song_id=song.id)); db.commit()
    return {"message": "Added to history"}

@router.get("", response_model=List[HistoryResponse])
def get_history(current_user: User = Depends(get_current_user), db: Session = Depends(get_db)):
    return db.query(History).filter(History.user_id == current_user.id).order_by(
        History.played_at.desc()).limit(50).all()

@router.delete("")
def clear_history(current_user: User = Depends(get_current_user), db: Session = Depends(get_db)):
    try:
        count = db.query(History).filter(History.user_id == current_user.id).delete()
        db.commit()
        return {"message": "History cleared", "deleted_count": count}
    except Exception as e:
        db.rollback()
        raise HTTPException(500, f"Failed to clear history: {str(e)}")''', 'python')
para('Every time a user plays a song, the frontend calls POST /history. The backend uses "get or create" for the song record, then creates a new History entry with the current UTC timestamp. History is returned ordered by played_at descending (newest first), limited to 50 entries. clear_history uses SQLAlchemy bulk delete with rollback on failure.')

# ── stats router ──
h('5.1.5 Statistics Router — backend/app/routers/stats.py', 3)
code('''from fastapi import APIRouter, Depends
from sqlalchemy.orm import Session
from sqlalchemy import func, desc
import random
from ..database import get_db
from ..models import Song, History, Playlist, Like, User
from ..auth import get_current_user

router = APIRouter(prefix="/stats", tags=["Statistics"])

@router.get("/trending")
def get_trending_songs(db: Session = Depends(get_db)):
    """Most played songs across all users, ranked by play count."""
    trending = db.query(Song, func.count(History.id).label(\'play_count\')).join(
        History, Song.id == History.song_id).group_by(Song.id).order_by(
        desc(\'play_count\')).limit(20).all()
    return {"trending": [{"youtube_video_id": s.youtube_video_id, "title": s.title,
                           "thumbnail": s.thumbnail, "channel": s.channel,
                           "play_count": pc} for s, pc in trending]}

@router.get("/recommendations")
def get_recommendations(current_user: User = Depends(get_current_user),
                        db: Session = Depends(get_db)):
    """3-strategy recommendation: liked artists → trending → database fallback."""
    recommendations, message, source = [], "", ""
    liked_songs = db.query(Song).join(Like, Song.id == Like.song_id).filter(
        Like.user_id == current_user.id).limit(10).all()
    if liked_songs:
        liked_artists  = set(s.channel for s in liked_songs if s.channel)
        liked_song_ids = set(s.id for s in liked_songs)
        for artist in liked_artists:
            recommendations.extend(db.query(Song).filter(
                Song.channel == artist, Song.id.notin_(liked_song_ids)).limit(10).all())
        message, source = "Based on artists you liked", "liked_artists"
    if len(recommendations) < 50:
        existing = set(s.id for s in recommendations)
        for song, _ in db.query(Song, func.count(History.id)).join(
                History, Song.id == History.song_id).group_by(Song.id).order_by(
                desc(func.count(History.id))).limit(50).all():
            if song.id not in existing and len(recommendations) < 50:
                recommendations.append(song)
        if not liked_songs: message, source = "Trending songs", "trending"
    if len(recommendations) < 50:
        existing = set(s.id for s in recommendations)
        recommendations.extend(db.query(Song).filter(
            Song.id.notin_(existing)).limit(50).all())
        if not message: message, source = "Songs from your library", "random"
    random.shuffle(recommendations)
    return {"recommendations": [{"youtube_video_id": s.youtube_video_id, "title": s.title,
                                  "thumbnail": s.thumbnail, "channel": s.channel}
                                 for s in recommendations[:50]],
            "source": source, "message": message, "total": len(recommendations[:50])}

@router.get("/analytics")
def get_analytics(current_user: User = Depends(get_current_user), db: Session = Depends(get_db)):
    return {
        "total_songs_played": db.query(func.count(History.id)).filter(
            History.user_id == current_user.id).scalar() or 0,
        "total_playlists": db.query(func.count(Playlist.id)).filter(
            Playlist.user_id == current_user.id).scalar() or 0,
        "total_likes": db.query(func.count(Like.id)).filter(
            Like.user_id == current_user.id).scalar() or 0,
    }

@router.post("/like/{youtube_video_id}")
def like_song(youtube_video_id: str, current_user: User = Depends(get_current_user),
              db: Session = Depends(get_db)):
    song = db.query(Song).filter(Song.youtube_video_id == youtube_video_id).first()
    if not song: return {"error": "Song not found"}
    if db.query(Like).filter(Like.user_id == current_user.id,
                              Like.song_id == song.id).first():
        return {"message": "Already liked"}
    db.add(Like(user_id=current_user.id, song_id=song.id)); db.commit()
    return {"message": "Song liked"}

@router.delete("/like/{youtube_video_id}")
def unlike_song(youtube_video_id: str, current_user: User = Depends(get_current_user),
                db: Session = Depends(get_db)):
    song = db.query(Song).filter(Song.youtube_video_id == youtube_video_id).first()
    if not song: return {"error": "Song not found"}
    like = db.query(Like).filter(Like.user_id == current_user.id,
                                  Like.song_id == song.id).first()
    if like: db.delete(like); db.commit(); return {"message": "Song unliked"}
    return {"message": "Not liked"}

@router.get("/liked")
def get_liked_songs(current_user: User = Depends(get_current_user), db: Session = Depends(get_db)):
    likes = db.query(Like, Song).join(Song, Like.song_id == Song.id).filter(
        Like.user_id == current_user.id).order_by(desc(Like.created_at)).all()
    return {"liked_songs": [{"youtube_video_id": s.youtube_video_id, "title": s.title,
                              "thumbnail": s.thumbnail, "channel": s.channel,
                              "liked_at": l.created_at.isoformat()} for l, s in likes]}''', 'python')
para('get_trending_songs uses SQLAlchemy func.count() with JOIN and GROUP BY to count history entries per song across all users. get_recommendations implements 3-strategy fallback. get_analytics uses three scalar count queries. Like/unlike is idempotent — liking an already-liked song returns "Already liked" without creating a duplicate.')

# ── offline router ──
h('5.1.6 Offline Router — backend/app/routers/offline.py', 3)
code('''from fastapi import APIRouter, Depends, HTTPException, status
from sqlalchemy.orm import Session
from sqlalchemy import desc
from ..database import get_db
from ..models import OfflineSong, Song, User
from ..auth import get_current_user
from ..schemas import HistoryCreate

router = APIRouter(prefix="/offline", tags=["Offline"])

@router.post("/save/{youtube_video_id}", status_code=201)
async def save_song_offline(youtube_video_id: str, song_data: HistoryCreate,
                             current_user: User = Depends(get_current_user),
                             db: Session = Depends(get_db)):
    song = db.query(Song).filter(Song.youtube_video_id == youtube_video_id).first()
    if not song:
        song = Song(youtube_video_id=song_data.youtube_video_id, title=song_data.title,
                    thumbnail=song_data.thumbnail, channel=song_data.channel)
        db.add(song); db.commit(); db.refresh(song)
    if db.query(OfflineSong).filter(OfflineSong.user_id == current_user.id,
                                     OfflineSong.song_id == song.id).first():
        raise HTTPException(400, "Song already saved for offline")
    offline = OfflineSong(user_id=current_user.id, song_id=song.id)
    db.add(offline); db.commit(); db.refresh(offline)
    return {"message": "Song saved for offline",
            "song": {"youtube_video_id": song.youtube_video_id, "title": song.title,
                     "thumbnail": song.thumbnail, "channel": song.channel},
            "saved_at": offline.saved_at}

@router.get("")
async def get_offline_songs(current_user: User = Depends(get_current_user),
                             db: Session = Depends(get_db)):
    rows = db.query(OfflineSong, Song).join(Song, OfflineSong.song_id == Song.id).filter(
        OfflineSong.user_id == current_user.id).order_by(desc(OfflineSong.saved_at)).all()
    return {"offline_songs": [{"id": o.id, "saved_at": o.saved_at,
                                "song": {"youtube_video_id": s.youtube_video_id,
                                         "title": s.title, "thumbnail": s.thumbnail,
                                         "channel": s.channel}} for o, s in rows],
            "total": len(rows)}

@router.delete("/{youtube_video_id}")
async def remove_offline_song(youtube_video_id: str,
                               current_user: User = Depends(get_current_user),
                               db: Session = Depends(get_db)):
    song = db.query(Song).filter(Song.youtube_video_id == youtube_video_id).first()
    if not song: raise HTTPException(404, "Song not found")
    offline = db.query(OfflineSong).filter(OfflineSong.user_id == current_user.id,
                                            OfflineSong.song_id == song.id).first()
    if not offline: raise HTTPException(404, "Song not saved for offline")
    db.delete(offline); db.commit()
    return {"message": "Song removed from offline"}

@router.get("/check/{youtube_video_id}")
async def check_offline_status(youtube_video_id: str,
                                current_user: User = Depends(get_current_user),
                                db: Session = Depends(get_db)):
    song = db.query(Song).filter(Song.youtube_video_id == youtube_video_id).first()
    if not song: return {"is_saved": False}
    offline = db.query(OfflineSong).filter(OfflineSong.user_id == current_user.id,
                                            OfflineSong.song_id == song.id).first()
    return {"is_saved": offline is not None,
            "saved_at": offline.saved_at if offline else None}''', 'python')
para('The offline feature is a bookmarking system — it saves song metadata to the database so users can access their saved songs list. check_offline_status is used by the frontend to show the correct icon (checkmark if saved, floppy disk if not) on song cards.')
pb()

# ── stats router ──
h('5.1.5 Stats Router — backend/app/routers/stats.py', 3)
code('''from fastapi import APIRouter, Depends
from sqlalchemy.orm import Session
from sqlalchemy import func
from ..database import get_db
from ..models import Song, History, Like, User
from ..auth import get_current_user

router = APIRouter(prefix="/stats", tags=["Stats"])

@router.post("/like/{video_id}", status_code=201)
def like_song(video_id: str, current_user: User = Depends(get_current_user),
              db: Session = Depends(get_db)):
    song = db.query(Song).filter(Song.youtube_video_id == video_id).first()
    if not song: raise HTTPException(404, "Song not found")
    if db.query(Like).filter(Like.user_id == current_user.id, Like.song_id == song.id).first():
        raise HTTPException(400, "Already liked")
    db.add(Like(user_id=current_user.id, song_id=song.id)); db.commit()
    return {"message": "Song liked"}

@router.delete("/like/{video_id}")
def unlike_song(video_id: str, current_user: User = Depends(get_current_user),
                db: Session = Depends(get_db)):
    song = db.query(Song).filter(Song.youtube_video_id == video_id).first()
    if not song: raise HTTPException(404, "Song not found")
    like = db.query(Like).filter(Like.user_id == current_user.id, Like.song_id == song.id).first()
    if not like: raise HTTPException(404, "Like not found")
    db.delete(like); db.commit()
    return {"message": "Song unliked"}

@router.get("/liked")
def get_liked_songs(current_user: User = Depends(get_current_user), db: Session = Depends(get_db)):
    likes = db.query(Like).filter(Like.user_id == current_user.id).all()
    return {"liked_songs": [{"youtube_video_id": l.song.youtube_video_id,
                              "title": l.song.title, "thumbnail": l.song.thumbnail,
                              "channel": l.song.channel} for l in likes]}

@router.get("/recommendations")
def get_recommendations(current_user: User = Depends(get_current_user),
                         db: Session = Depends(get_db)):
    liked = db.query(Like).filter(Like.user_id == current_user.id).limit(10).all()
    recs = []; seen = set()
    if liked:
        artists = list({l.song.channel for l in liked if l.song.channel})
        for artist in artists[:5]:
            songs = db.query(Song).filter(Song.channel.ilike(f"%{artist}%")).limit(10).all()
            for s in songs:
                if s.youtube_video_id not in seen:
                    recs.append(s); seen.add(s.youtube_video_id)
        msg = "Based on your liked songs"
    if len(recs) < 50:
        trending = db.query(Song, func.count(History.id).label("plays")).join(
            History, Song.id == History.song_id).group_by(Song.id).order_by(
            func.count(History.id).desc()).limit(50).all()
        for s, _ in trending:
            if s.youtube_video_id not in seen:
                recs.append(s); seen.add(s.youtube_video_id)
        msg = "Trending songs" if not liked else msg
    import random; random.shuffle(recs)
    return {"recommendations": [{"youtube_video_id": s.youtube_video_id, "title": s.title,
                                  "thumbnail": s.thumbnail, "channel": s.channel}
                                 for s in recs[:50]], "message": msg if liked else "Trending songs"}

@router.get("/trending")
def get_trending(db: Session = Depends(get_db)):
    trending = db.query(Song, func.count(History.id).label("play_count")).join(
        History, Song.id == History.song_id).group_by(Song.id).order_by(
        func.count(History.id).desc()).limit(20).all()
    return {"trending": [{"youtube_video_id": s.youtube_video_id, "title": s.title,
                           "thumbnail": s.thumbnail, "channel": s.channel,
                           "play_count": pc} for s, pc in trending]}

@router.get("/analytics")
def get_analytics(current_user: User = Depends(get_current_user), db: Session = Depends(get_db)):
    total_plays = db.query(History).filter(History.user_id == current_user.id).count()
    total_likes = db.query(Like).filter(Like.user_id == current_user.id).count()
    top_songs = db.query(Song, func.count(History.id).label("plays")).join(
        History, Song.id == History.song_id).filter(
        History.user_id == current_user.id).group_by(Song.id).order_by(
        func.count(History.id).desc()).limit(10).all()
    return {"total_plays": total_plays, "total_likes": total_likes,
            "top_songs": [{"title": s.title, "channel": s.channel,
                            "plays": p} for s, p in top_songs]}''', 'python')
para('The recommendations endpoint implements a 3-strategy fallback: liked artist songs first, then trending songs, then any database songs. random.shuffle() ensures variety on each call. The analytics endpoint uses SQLAlchemy func.count() with GROUP BY for efficient aggregation queries.')

# ── offline router ──
h('5.1.6 Offline Router — backend/app/routers/offline.py', 3)
code('''from fastapi import APIRouter, Depends, HTTPException, status
from sqlalchemy.orm import Session
from ..database import get_db
from ..models import Song, OfflineSong, User
from ..schemas import SongBase
from ..auth import get_current_user

router = APIRouter(prefix="/offline", tags=["Offline"])

@router.post("/save/{video_id}", status_code=201)
def save_offline(video_id: str, song_data: SongBase,
                 current_user: User = Depends(get_current_user),
                 db: Session = Depends(get_db)):
    song = db.query(Song).filter(Song.youtube_video_id == video_id).first()
    if not song:
        song = Song(youtube_video_id=video_id, title=song_data.title,
                    thumbnail=song_data.thumbnail, channel=song_data.channel)
        db.add(song); db.commit(); db.refresh(song)
    if db.query(OfflineSong).filter(OfflineSong.user_id == current_user.id,
                                     OfflineSong.song_id == song.id).first():
        raise HTTPException(400, "Already saved offline")
    db.add(OfflineSong(user_id=current_user.id, song_id=song.id)); db.commit()
    return {"message": "Song saved for offline"}

@router.get("")
def get_offline_songs(current_user: User = Depends(get_current_user),
                      db: Session = Depends(get_db)):
    offline = db.query(OfflineSong).filter(OfflineSong.user_id == current_user.id).all()
    return {"offline_songs": [{"song": {"youtube_video_id": o.song.youtube_video_id,
                                         "title": o.song.title, "thumbnail": o.song.thumbnail,
                                         "channel": o.song.channel},
                                "saved_at": o.saved_at} for o in offline]}

@router.delete("/{video_id}")
def remove_offline(video_id: str, current_user: User = Depends(get_current_user),
                   db: Session = Depends(get_db)):
    song = db.query(Song).filter(Song.youtube_video_id == video_id).first()
    if not song: raise HTTPException(404, "Song not found")
    offline = db.query(OfflineSong).filter(OfflineSong.user_id == current_user.id,
                                            OfflineSong.song_id == song.id).first()
    if not offline: raise HTTPException(404, "Song not saved offline")
    db.delete(offline); db.commit()
    return {"message": "Removed from offline"}''', 'python')
para('Offline songs are bookmarks — they store a reference to the song in the database. The 💾 icon in the UI toggles the saved state. The offline page displays all saved songs so users can quickly access their bookmarked content.')

h('5.2 Frontend Implementation', 2)

h('5.2.1 Application Root — frontend/src/App.jsx', 3)
para('App.jsx is the root component. It manages global state: currentSong, queue (array), playHistory (array used as stack for previous navigation), isAuthenticated, theme, and showFullScreen. It renders the Sidebar, page routes, FullScreenPlayer, MiniPlayer, Onboarding, and ToastContainer.')
code('''// Key state and handlers in App.jsx
const [currentSong, setCurrentSong] = useState(null);
const [queue, setQueue] = useState([]);
const [playHistory, setPlayHistory] = useState([]);

const handlePlaySong = (song) => {
  setCurrentSong(song); setShowFullScreen(true); setIsPlaying(true);
  setPlayHistory(prev => [...prev, song]);
};

const playNext = () => {
  if (queue.length > 0) {
    const nextSong = queue[0];
    setCurrentSong(nextSong); setQueue(prev => prev.slice(1));
    setPlayHistory(prev => [...prev, nextSong]);
  }
};

const playPrevious = () => {
  if (playHistory.length > 1) {
    const newHistory = [...playHistory];
    newHistory.pop(); // remove current
    setCurrentSong(newHistory[newHistory.length - 1]);
    setPlayHistory(newHistory);
  }
};

// Route structure
<Routes>
  <Route path="/"          element={<Home onPlaySong={handlePlaySong} />} />
  <Route path="/search"    element={<Search onPlaySong={handlePlaySong} />} />
  <Route path="/mood"      element={<MoodSlider onPlaySong={handlePlaySong} />} />
  <Route path="/playlists" element={<Playlists />} />
  <Route path="/history"   element={<History onPlaySong={handlePlaySong} />} />
  <Route path="/offline"   element={<Offline onPlaySong={handlePlaySong} />} />
  <Route path="/trending"  element={<Trending onPlaySong={handlePlaySong} />} />
  <Route path="/analytics" element={<Analytics />} />
</Routes>''', 'jsx')

h('5.2.2 Home Page — frontend/src/pages/Home.jsx', 3)
para('Home.jsx fetches 5 data sources in parallel using Promise.all: history, recommendations, liked songs, related artists, and trending songs. It displays recommendations in a grid of 16 songs (8 + artist section + 8), with a "Show More" button that increments displayCount by 16 each click. The 3-strategy recommendation engine is driven by the backend /stats/recommendations endpoint.')
code('''// Parallel data fetch on mount
const [historyRes, recsRes, likedRes, artistsRes, trendingRes] = await Promise.all([
  api.get('/history'),
  api.get('/stats/recommendations'),
  api.get('/stats/liked'),
  api.get('/stats/related-artists'),
  api.get('/stats/trending')
]);

// Display logic: show first 8, then artists, then next 8
{allRecommendations.slice(0, Math.min(8, displayCount)).map(song => <SongCard />)}
{relatedArtists.length > 0 && displayCount > 8 && <ArtistsSection />}
{displayCount > 8 && allRecommendations.slice(8, displayCount).map(song => <SongCard />)}

// Show More increments by 16
const showMoreSongs = () => setDisplayCount(prev => prev + 16);''', 'jsx')

h('5.2.3 Search Page — frontend/src/pages/Search.jsx', 3)
para('Search.jsx checks the 12-hour localStorage cache before making an API call. If a cache hit is found, results are displayed instantly with a toast notification. Voice search is triggered by the 🎤 button. Language filter appends the language name to the search query. Pagination uses YouTube\'s nextPageToken. The playlist modal allows creating new playlists or adding to existing ones inline.')
code('''// Cache-first search
const handleSearch = async (e, searchQuery = null) => {
  const cached = searchCache.get(queryToSearch, language);
  if (cached) {
    setResults(cached.results);
    window.showToast("Loaded from 12-hour cache - API saved!", "success");
    return;
  }
  // Fresh API call
  const response = await api.get(`/youtube/search?query=${encodeURIComponent(q)}&language=${language}`);
  setResults(response.data.results);
  searchCache.set(queryToSearch, language, { results: response.data.results,
                                              nextPageToken: response.data.nextPageToken });
};

// Load more via pageToken
const loadMoreResults = async () => {
  const response = await api.get(`/youtube/search?query=${q}&language=${lang}&pageToken=${nextPageToken}`);
  setResults(prev => [...prev, ...response.data.results]);
};''', 'jsx')

h('5.2.4 Mini Player — frontend/src/components/MiniPlayer.jsx', 3)
para('MiniPlayer renders at the bottom of the screen when a song is playing and the full-screen player is closed. It shows a progress bar, thumbnail, title, artist, and prev/play/next controls. Clicking the song info or the expand button opens the full-screen player.')
code('''function MiniPlayer({ currentSong, isPlaying, onPlayPause, onNext, onPrevious, onExpand, progress }) {
  if (!currentSong) return null;
  return (
    <div className="mini-player">
      <div className="mini-player-progress" style={{ width: `${progress}%` }}></div>
      <div className="mini-player-content">
        <div className="mini-player-song" onClick={onExpand}>
          <img src={currentSong.thumbnail} className="mini-player-thumbnail" />
          <div className="mini-player-info">
            <h4>{currentSong.title}</h4>
            <p>{currentSong.channel}</p>
          </div>
        </div>
        <div className="mini-player-controls">
          <button onClick={onPrevious}>⏮</button>
          <button onClick={onPlayPause}>{isPlaying ? "⏸" : "▶"}</button>
          <button onClick={onNext}>⏭</button>
        </div>
        <button onClick={onExpand}>⬆</button>
      </div>
    </div>
  );
}''', 'jsx')

h('5.2.5 Karaoke Mode — frontend/src/components/KaraokeMode.jsx', 3)
para('KaraokeMode fetches lyrics from the /ai/lyrics/:title endpoint. A setInterval running every 500ms reads the YouTube player\'s getCurrentTime() and finds the matching lyric line. The active line is highlighted and auto-scrolled into view using scrollIntoView({behavior: "smooth", block: "center"}). Press ESC to exit.')
code('''// Auto-scroll and highlight based on playback time
useEffect(() => {
  if (!player || !lyrics.length) return;
  const interval = setInterval(() => {
    const currentTime = player.getCurrentTime();
    let newIndex = 0;
    for (let i = 0; i < lyrics.length; i++) {
      if (lyrics[i].time <= currentTime) newIndex = i;
      else break;
    }
    if (newIndex !== currentLineIndex) {
      setCurrentLineIndex(newIndex);
      lyricsContainerRef.current?.children[newIndex]?.scrollIntoView({
        behavior: "smooth", block: "center"
      });
    }
  }, 500);
  return () => clearInterval(interval);
}, [player, lyrics, currentLineIndex]);''', 'jsx')

h('5.2.6 Voice Search — frontend/src/components/VoiceSearch.jsx', 3)
para('VoiceSearch uses the Web Speech API (SpeechRecognition / webkitSpeechRecognition). It starts listening automatically on mount. Both interim (real-time) and final transcripts are displayed. The processVoiceCommand() function strips common prefixes like "play", "search for", "find" before passing the query to the search handler.')
code('''const processVoiceCommand = (command) => {
  const prefixes = ["hey play","play","search for","find","show me",
                    "i want to listen to","play me","can you play","please play"];
  let processed = command.toLowerCase().trim();
  for (const prefix of prefixes) {
    if (processed.startsWith(prefix)) {
      processed = processed.substring(prefix.length).trim();
      break;
    }
  }
  return processed;
};
// Example: "play Arijit Singh" -> "Arijit Singh"
// Example: "search for Hindi songs" -> "Hindi songs"''', 'jsx')

h('5.2.7 Social Share — frontend/src/components/SocialShare.jsx', 3)
para('SocialShare generates platform-specific share URLs using the YouTube video URL and song metadata. It supports 6 platforms plus clipboard copy and native Web Share API (mobile). Each platform button opens a 600x400 popup window.')
code('''const shareUrl  = `https://www.youtube.com/watch?v=${song.youtube_video_id || song.videoId}`;
const shareText = `Check out: ${song.title} by ${song.channel || song.channelTitle}`;
const shareLinks = {
  whatsapp: `https://wa.me/?text=${encodeURIComponent(shareText + " " + shareUrl)}`,
  facebook: `https://www.facebook.com/sharer/sharer.php?u=${encodeURIComponent(shareUrl)}`,
  twitter:  `https://twitter.com/intent/tweet?text=${encodeURIComponent(shareText)}&url=${encodeURIComponent(shareUrl)}`,
  telegram: `https://t.me/share/url?url=${encodeURIComponent(shareUrl)}&text=${encodeURIComponent(shareText)}`,
  linkedin: `https://www.linkedin.com/sharing/share-offsite/?url=${encodeURIComponent(shareUrl)}`,
  reddit:   `https://reddit.com/submit?url=${encodeURIComponent(shareUrl)}&title=${encodeURIComponent(song.title)}`
};
const handleShare = (platform) =>
  window.open(shareLinks[platform], "_blank", "width=600,height=400");''', 'jsx')

h('5.2.8 Onboarding — frontend/src/components/Onboarding.jsx', 3)
para('Onboarding shows a 6-step modal for first-time users. It checks localStorage for the "music_sagar_onboarding_completed" key. If not set, the modal appears. Each step shows an icon, title, description, and feature list. Progress dots at the top indicate current step. On completion or skip, the key is set in localStorage so it never shows again.')
code('''const steps = [
  { title: "Welcome to Music Sagar!", icon: "🎉",
    features: ["Stream unlimited music","Create playlists","Discover songs","Track history"] },
  { title: "Search & Discover",       icon: "🔎",
    features: ["Search by song/artist","Filter by language","Instant YouTube results"] },
  { title: "Mood-Based Discovery",    icon: "😊",
    features: ["Mood slider (Sad to Energetic)","Personalized recommendations"] },
  { title: "Playlists & History",     icon: "📝",
    features: ["Create unlimited playlists","View listening history","Like songs"] },
  { title: "Keyboard Shortcuts",      icon: "⚡",
    features: ["Space: Play/Pause","K: Karaoke Mode","Ctrl+Q: Add to queue"] },
  { title: "You are All Set!",        icon: "✨",
    features: ["Browse recommendations","Search songs","Create your first playlist"] }
];
const handleComplete = () => {
  localStorage.setItem("music_sagar_onboarding_completed", "true");
  setShow(false);
};''', 'jsx')

h('5.2.9 Sidebar — frontend/src/components/Sidebar.jsx', 3)
para('Sidebar uses React Router\'s useLocation hook to detect the active route and apply the "active" CSS class. It renders 8 navigation links, a theme toggle button, and a logout button. The logo uses an inline SVG treble clef design.')
code('''function Sidebar({ onThemeToggle, currentTheme, username, onLogout }) {
  const location = useLocation();
  const isActive = (path) => location.pathname === path ? "active" : "";
  return (
    <div className="sidebar">
      <div className="logo"><span>Music Sagar</span></div>
      <nav className="nav-menu">
        <Link to="/"          className={`nav-item ${isActive("/")}`}>🏠 Home</Link>
        <Link to="/search"    className={`nav-item ${isActive("/search")}`}>🔍 Search</Link>
        <Link to="/mood"      className={`nav-item ${isActive("/mood")}`}>🎭 Mood</Link>
        <Link to="/playlists" className={`nav-item ${isActive("/playlists")}`}>📚 Playlists</Link>
        <Link to="/offline"   className={`nav-item ${isActive("/offline")}`}>📥 Offline</Link>
        <Link to="/history"   className={`nav-item ${isActive("/history")}`}>🕒 History</Link>
        <Link to="/trending"  className={`nav-item ${isActive("/trending")}`}>🔥 Trending</Link>
        <Link to="/analytics" className={`nav-item ${isActive("/analytics")}`}>📊 Analytics</Link>
      </nav>
      <div className="sidebar-footer">
        <button onClick={onThemeToggle}>{currentTheme === "dark" ? "☀️" : "🌙"}</button>
        <button onClick={onLogout}>🚪 Logout</button>
      </div>
    </div>
  );
}''', 'jsx')

h('5.3 Testing Results', 2)
table(
    ['Test ID', 'Test Case', 'Input', 'Expected', 'Result', 'Status'],
    [
        ['TC1',  'Valid user signup',              'New username/email/password',    '201 Created',          '201 Created',          'PASS'],
        ['TC2',  'Duplicate email signup',         'Existing email',                 '400 Bad Request',      '400 Bad Request',      'PASS'],
        ['TC3',  'Valid login',                    'Correct credentials',            'JWT token returned',   'JWT token returned',   'PASS'],
        ['TC4',  'Wrong password login',           'Incorrect password',             '401 Unauthorized',     '401 Unauthorized',     'PASS'],
        ['TC5',  'Protected route no token',       'No Authorization header',        '401 Unauthorized',     '401 Unauthorized',     'PASS'],
        ['TC6',  'YouTube search',                 '"Arijit Singh"',                 '20 results returned',  '20 results returned',  'PASS'],
        ['TC7',  'Language filter search',         'query="songs", lang="hindi"',    'Hindi songs returned', 'Hindi songs returned', 'PASS'],
        ['TC8',  'Cache hit (repeat search)',      'Same query within 12 hours',     'Instant cache result', 'Instant cache result', 'PASS'],
        ['TC9',  'Create playlist',                'name="My Playlist"',             '201 Created',          '201 Created',          'PASS'],
        ['TC10', 'Add song to playlist',           'Valid song data',                'Song added 200 OK',    'Song added 200 OK',    'PASS'],
        ['TC11', 'Duplicate song in playlist',     'Same song twice',                '400 Bad Request',      '400 Bad Request',      'PASS'],
        ['TC12', 'Delete playlist',                'Valid playlist ID',              '204 No Content',       '204 No Content',       'PASS'],
        ['TC13', 'Access other user playlist',     'Different user playlist ID',     '404 Not Found',        '404 Not Found',        'PASS'],
        ['TC14', 'Like a song',                    'Valid video ID',                 '201 Created',          '201 Created',          'PASS'],
        ['TC15', 'Save song offline',              'Valid song data',                '201 Created',          '201 Created',          'PASS'],
        ['TC16', 'Clear history',                  'Authenticated user',             'History cleared',      'History cleared',      'PASS'],
        ['TC17', 'Voice search NLP processing',   '"play Arijit Singh"',            '"Arijit Singh"',       '"Arijit Singh"',       'PASS'],
        ['TC18', 'Social share URL generation',   'Song with video ID',             'Valid share URLs',     'Valid share URLs',     'PASS'],
        ['TC19', 'API key rotation on 403',        'Quota exceeded key',             'Next key used',        'Next key used',        'PASS'],
        ['TC20', 'Mini player visibility',         'Song playing, player closed',    'Mini player shown',    'Mini player shown',    'PASS'],
    ]
)
pb()

# ─────────────────────────────────────────────────────────────────────────────
# CHAPTER 6
# ─────────────────────────────────────────────────────────────────────────────
h('CHAPTER 6: RESULTS AND DISCUSSION', 1)

h('6.1 Performance Metrics', 2)
table(
    ['Metric', 'Measured Value', 'Target', 'Status'],
    [
        ['Fresh search response time',     '1.2 – 2.8 seconds',  '< 5 seconds',  'ACHIEVED'],
        ['Cached search response time',    '< 50 milliseconds',  '< 100ms',      'ACHIEVED'],
        ['Login / JWT generation',         '< 300 milliseconds', '< 500ms',      'ACHIEVED'],
        ['Playlist creation',              '< 200 milliseconds', '< 500ms',      'ACHIEVED'],
        ['History fetch (50 records)',     '< 150 milliseconds', '< 200ms',      'ACHIEVED'],
        ['Analytics query',                '< 200 milliseconds', '< 500ms',      'ACHIEVED'],
        ['API cache hit rate (repeat)',    '~80% reduction',     '> 70%',        'ACHIEVED'],
        ['Daily quota capacity',           '300,000 units/day',  '> 100,000',    'ACHIEVED'],
        ['Concurrent API key failover',    '< 100ms retry',      'Seamless',     'ACHIEVED'],
    ]
)
img(r'C:\Users\hp\Documents\a\b\Performance_Metrics.png', 'Figure 6.1: Performance Metrics — Actual vs Target', width=6.2)

h('6.2 Feature Completeness', 2)
table(
    ['Feature', 'Status', 'Notes'],
    [
        ['User Registration & Login',          'Complete', 'JWT + bcrypt, 7-day token'],
        ['YouTube Music Search',               'Complete', '20 results, language filter, pagination'],
        ['30-Key API Rotation',                'Complete', '300,000 daily quota units'],
        ['12-Hour Search Cache',               'Complete', 'localStorage, 200 entries, LRU eviction'],
        ['Playlist CRUD',                      'Complete', 'Create, rename, delete, add/remove songs'],
        ['Listening History',                  'Complete', 'Auto-track, view last 50, clear all'],
        ['Like / Unlike Songs',                'Complete', 'Toggle, persisted per user'],
        ['Offline Bookmarking',                'Complete', 'Save/remove, 💾 icon toggle'],
        ['Trending Songs',                     'Complete', 'By play count across all users'],
        ['Mood-Based Discovery',               'Complete', '0-100 slider, debounced API calls'],
        ['Personalized Recommendations',       'Complete', '3-strategy: liked artists → trending → DB'],
        ['Analytics Dashboard',               'Complete', 'Total plays, likes, top songs, top artists'],
        ['Voice Search',                       'Complete', 'Web Speech API + NLP prefix stripping'],
        ['Social Sharing',                     'Complete', '6 platforms + clipboard + native share'],
        ['Karaoke Mode',                       'Complete', 'Synchronized lyrics, auto-scroll'],
        ['Mini Player',                        'Complete', 'Progress bar, prev/play/next, expand'],
        ['Dark / Light Theme',                 'Complete', 'CSS variables, persisted preference'],
        ['Onboarding Flow',                    'Complete', '6-step modal, localStorage completion flag'],
        ['Related Artists Section',            'Complete', 'Derived from liked song channels'],
        ['Queue Management',                   'Complete', 'Add, remove, clear, next/previous'],
    ]
)
img(r'C:\Users\hp\Documents\a\b\Feature_Completeness.png', 'Figure 6.2: Feature Completeness — All 20 Features at 100%', width=6.2)

h('6.3 Challenges and Solutions', 2)
table(
    ['Challenge', 'Solution'],
    [
        ['YouTube API quota exhaustion (10,000 units/key/day)',
         'Implemented 30-key rotation with failed_keys set; 12-hour cache reduces consumption ~80%'],
        ['Concurrent requests hitting same exhausted key',
         'Module-level failed_keys set persists across requests; immediate retry with next key on 403'],
        ['Duplicate songs across users wasting database space',
         'Songs table uses youtube_video_id as unique index; "get or create" pattern reuses existing records'],
        ['Search results blocking on slow duration fetch',
         'Duration fetch runs with asyncio.wait_for(timeout=3.0); search results returned even if durations fail'],
        ['Previous song navigation without server state',
         'Client-side playHistory array acts as a stack; pop() removes current, last element is previous'],
        ['Voice search browser compatibility',
         'webkitSpeechRecognition fallback for Safari; unsupported browsers show clear error message'],
        ['Cache growing too large in localStorage',
         'MAX_CACHE_SIZE=200 with LRU eviction; clearExpired() runs every 2 hours automatically'],
        ['JWT token expiry handling',
         'Axios interceptor attaches token; 401 responses redirect to login page'],
    ]
)
pb()

# ─────────────────────────────────────────────────────────────────────────────
# CHAPTER 7
# ─────────────────────────────────────────────────────────────────────────────
h('CHAPTER 7: CONCLUSIONS AND FUTURE SCOPE', 1)

h('7.1 Conclusions', 2)
para('The Music Streaming System successfully demonstrates the design and development of a production-quality full-stack web application. The project achieves all stated objectives:')
numbered_reset()
for item in [
    'A secure multi-user system with JWT authentication and bcrypt password hashing was implemented and tested.',
    'The YouTube Data API v3 integration with 30-key rotation provides 300,000 daily quota units with automatic failover.',
    'The 12-hour client-side search cache reduces API consumption by approximately 80% for repeat searches.',
    'Complete playlist management, listening history, likes, offline bookmarking, and analytics are fully functional.',
    'Advanced features including voice search, social sharing, karaoke mode, mood-based discovery, and mini player enhance the user experience significantly.',
    'All 20 test cases passed, confirming the correctness of authentication, data isolation, API integration, and caching.',
    'The three-tier architecture (React.js + FastAPI + SQLite) provides a clean separation of concerns and a maintainable codebase.',
]:
    numbered(item)
para('The project demonstrates practical application of modern web development concepts including REST API design, ORM-based database management, stateless authentication, client-side caching, asynchronous programming, and component-based UI development.')

h('7.2 Future Scope', 2)
numbered_reset()
for item in [
    'Real-time lyrics synchronization using a dedicated lyrics API (Musixmatch, Genius) instead of AI-generated placeholders.',
    'Collaborative playlists allowing multiple users to add songs to a shared playlist.',
    'Music recommendations using machine learning — collaborative filtering based on listening patterns of similar users.',
    'Push notifications for new trending songs or playlist updates using Web Push API.',
    'Progressive Web App (PWA) support with service workers for true offline playback caching.',
    'Social features: follow other users, see their public playlists, share listening activity.',
    'Audio equalizer and playback speed controls using the Web Audio API.',
    'Migration from SQLite to PostgreSQL for production deployment with connection pooling.',
    'Mobile application using React Native sharing the same FastAPI backend.',
    'Admin dashboard for monitoring API quota usage, user statistics, and system health.',
]:
    numbered(item)
pb()

# ─────────────────────────────────────────────────────────────────────────────
# REFERENCES
# ─────────────────────────────────────────────────────────────────────────────
h('REFERENCES', 1)
for ref in [
    '[1] Tiangolo, S. (2023). FastAPI Documentation. https://fastapi.tiangolo.com/',
    '[2] Meta Platforms. (2023). React Documentation. https://react.dev/',
    '[3] SQLAlchemy Authors. (2023). SQLAlchemy ORM Documentation. https://docs.sqlalchemy.org/',
    '[4] Google LLC. (2023). YouTube Data API v3 Reference. https://developers.google.com/youtube/v3',
    '[5] Auth0. (2023). JSON Web Tokens Introduction. https://jwt.io/introduction',
    '[6] Python Software Foundation. (2023). Python 3 Documentation. https://docs.python.org/3/',
    '[7] Mozilla Developer Network. (2023). Web Speech API. https://developer.mozilla.org/en-US/docs/Web/API/Web_Speech_API',
    '[8] Pydantic Authors. (2023). Pydantic v2 Documentation. https://docs.pydantic.dev/',
    '[9] Encode. (2023). HTTPX Documentation. https://www.python-httpx.org/',
    '[10] Vite Authors. (2023). Vite Build Tool Documentation. https://vitejs.dev/',
    '[11] React Router Authors. (2023). React Router v6 Documentation. https://reactrouter.com/',
    '[12] Axios Authors. (2023). Axios HTTP Client Documentation. https://axios-http.com/',
    '[13] Passlib Authors. (2023). Passlib Documentation. https://passlib.readthedocs.io/',
    '[14] python-docx Authors. (2023). python-docx Documentation. https://python-docx.readthedocs.io/',
    '[15] MDN Web Docs. (2023). Web Storage API. https://developer.mozilla.org/en-US/docs/Web/API/Web_Storage_API',
]:
    para(ref, size=10)
pb()

# ─────────────────────────────────────────────────────────────────────────────
# GLOSSARY
# ─────────────────────────────────────────────────────────────────────────────
h('GLOSSARY', 1)
table(
    ['Term', 'Definition'],
    [
        ['API',           'Application Programming Interface — a set of rules for how software components communicate'],
        ['ASGI',          'Asynchronous Server Gateway Interface — Python async web server standard used by FastAPI/uvicorn'],
        ['bcrypt',        'Adaptive password hashing algorithm designed to be computationally expensive to resist brute-force'],
        ['CORS',          'Cross-Origin Resource Sharing — HTTP mechanism allowing browsers to make requests to different domains'],
        ['CRUD',          'Create, Read, Update, Delete — the four basic database operations'],
        ['FastAPI',       'Modern Python web framework for building REST APIs with automatic OpenAPI documentation'],
        ['JWT',           'JSON Web Token — compact, URL-safe token for stateless authentication'],
        ['LRU',           'Least Recently Used — cache eviction policy that removes the oldest accessed entry'],
        ['NLP',           'Natural Language Processing — processing and understanding human language text'],
        ['ORM',           'Object Relational Mapper — maps database tables to Python classes (SQLAlchemy)'],
        ['Pydantic',      'Python data validation library using type annotations; used by FastAPI for request/response schemas'],
        ['Quota',         'YouTube API usage limit — 10,000 units per key per day; search costs 100 units'],
        ['React',         'JavaScript library for building component-based user interfaces'],
        ['REST',          'Representational State Transfer — architectural style for web APIs using HTTP methods'],
        ['SPA',           'Single Page Application — web app that loads once and updates dynamically without full page reloads'],
        ['SQLAlchemy',    'Python SQL toolkit and ORM for database interaction'],
        ['SQLite',        'Lightweight file-based relational database; entire database stored in one .db file'],
        ['Token',         'In authentication context: a signed string proving identity without server-side session storage'],
        ['uvicorn',       'Lightning-fast ASGI server for running FastAPI applications'],
        ['Virtual DOM',   'React\'s in-memory representation of the DOM; enables efficient UI updates'],
        ['Web Speech API','Browser-native JavaScript API for speech recognition and synthesis'],
    ]
)
pb()

# ─────────────────────────────────────────────────────────────────────────────
# APPENDIX A
# ─────────────────────────────────────────────────────────────────────────────
h('APPENDIX A: PROJECT SETUP GUIDE', 1)

h('A.1 Prerequisites', 2)
for item in ['Python 3.10+ (py command on Windows)', 'Node.js 18+ and npm',
             'Git', 'YouTube Data API v3 keys (30 keys recommended)']:
    bullet(item)

h('A.2 Backend Setup', 2)
code('''# 1. Navigate to backend directory
cd backend

# 2. Create virtual environment
py -m venv venv

# 3. Activate virtual environment (Windows)
venv\\Scripts\\activate

# 4. Install dependencies
pip install -r requirements.txt

# 5. Configure environment variables
# Edit backend/.env and add your YouTube API keys:
# YOUTUBE_API_KEY_1=AIza...
# YOUTUBE_API_KEY_2=AIza...
# ... up to YOUTUBE_API_KEY_30

# 6. Initialize database
py init_fresh_db.py

# 7. Start backend server
uvicorn app.main:app --reload --port 8000''', 'bash')

h('A.3 Frontend Setup', 2)
code('''# 1. Navigate to frontend directory
cd frontend

# 2. Install dependencies
npm install

# 3. Start development server
npm run dev

# Frontend runs at: http://localhost:5173
# Backend API at:   http://localhost:8000
# API Docs at:      http://localhost:8000/docs''', 'bash')

h('A.4 Requirements File — backend/requirements.txt', 2)
code('''fastapi==0.104.1
uvicorn[standard]==0.24.0
sqlalchemy==2.0.23
python-jose[cryptography]==3.3.0
passlib[bcrypt]==1.7.4
python-dotenv==1.0.0
httpx==0.25.2
pydantic[email]==2.5.0
python-multipart==0.0.6''', 'text')
pb()

# ─────────────────────────────────────────────────────────────────────────────
# APPENDIX B
# ─────────────────────────────────────────────────────────────────────────────
h('APPENDIX B: API ENDPOINTS REFERENCE', 1)

h('B.1 Authentication Endpoints', 2)
table(
    ['Method', 'Endpoint', 'Auth Required', 'Description'],
    [
        ['POST', '/auth/signup',  'No',  'Register new user (username, email, password)'],
        ['POST', '/auth/login',   'No',  'Login and receive JWT token'],
        ['POST', '/auth/logout',  'No',  'Client-side logout (informational)'],
        ['GET',  '/auth/me',      'Yes', 'Get current authenticated user info'],
        ['POST', '/auth/theme',   'Yes', 'Update user theme preference (dark/light)'],
    ]
)

h('B.2 YouTube Endpoints', 2)
table(
    ['Method', 'Endpoint', 'Auth Required', 'Description'],
    [
        ['GET', '/youtube/search?query=&language=&pageToken=', 'No', 'Search YouTube music with optional language filter and pagination'],
    ]
)

h('B.3 Playlist Endpoints', 2)
table(
    ['Method', 'Endpoint', 'Auth Required', 'Description'],
    [
        ['POST',   '/playlists',                    'Yes', 'Create new playlist'],
        ['GET',    '/playlists',                    'Yes', 'Get all playlists for current user'],
        ['GET',    '/playlists/{id}/songs',         'Yes', 'Get all songs in a playlist'],
        ['POST',   '/playlists/{id}/add-song',      'Yes', 'Add song to playlist'],
        ['DELETE', '/playlists/{id}',               'Yes', 'Delete playlist'],
        ['PUT',    '/playlists/{id}',               'Yes', 'Rename playlist'],
    ]
)

h('B.4 History Endpoints', 2)
table(
    ['Method', 'Endpoint', 'Auth Required', 'Description'],
    [
        ['POST',   '/history', 'Yes', 'Add song to listening history'],
        ['GET',    '/history', 'Yes', 'Get last 50 history entries'],
        ['DELETE', '/history', 'Yes', 'Clear all listening history'],
    ]
)

h('B.5 Stats Endpoints', 2)
table(
    ['Method', 'Endpoint', 'Auth Required', 'Description'],
    [
        ['POST',   '/stats/like/{video_id}',      'Yes', 'Like a song'],
        ['DELETE', '/stats/like/{video_id}',      'Yes', 'Unlike a song'],
        ['GET',    '/stats/liked',                'Yes', 'Get all liked songs'],
        ['GET',    '/stats/recommendations',      'Yes', 'Get personalized recommendations'],
        ['GET',    '/stats/trending',             'No',  'Get trending songs by play count'],
        ['GET',    '/stats/analytics',            'Yes', 'Get user analytics data'],
        ['GET',    '/stats/related-artists',      'Yes', 'Get related artists from liked songs'],
    ]
)

h('B.6 Offline Endpoints', 2)
table(
    ['Method', 'Endpoint', 'Auth Required', 'Description'],
    [
        ['POST',   '/offline/save/{video_id}', 'Yes', 'Save song for offline access'],
        ['GET',    '/offline',                 'Yes', 'Get all offline saved songs'],
        ['DELETE', '/offline/{video_id}',      'Yes', 'Remove song from offline'],
    ]
)

# ─────────────────────────────────────────────────────────────────────────────
# SAVE
# ─────────────────────────────────────────────────────────────────────────────
doc.save('MUSIC_STREAMING_SYSTEM_BLACK_BOOK_v5.docx')
print("SUCCESS: MUSIC_STREAMING_SYSTEM_BLACK_BOOK_v5.docx generated!")
print(f"Sections: Title, Certificate, Abstract, Acknowledgement, TOC, Ch1-7, References, Glossary, Appendix A-B")
